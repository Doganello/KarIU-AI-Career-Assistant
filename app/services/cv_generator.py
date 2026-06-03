from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os
import re


class CVGenerator:
    def generate(self, graduate, lang="ru", vacancy="") -> str:
        """Генерация резюме в формате DOCX"""

        # Создаём документ
        doc = Document()

        # Настройка стилей
        self._setup_styles(doc)

        # Получаем фамилию для имени файла
        last_name = graduate.last_name if graduate.last_name else "выпускник"
        # Очищаем от спецсимволов для имени файла
        clean_last_name = re.sub(r'[^\w\s]', '', last_name).strip()

        # Добавляем заголовок
        title = doc.add_heading('РЕЗЮМЕ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Личная информация
        doc.add_heading('Личная информация', level=1)
        full_name = graduate.full_name if hasattr(graduate,
                                                  'full_name') else f"{graduate.last_name} {graduate.first_name}"
        doc.add_paragraph(f'ФИО: {full_name}')
        if graduate.birth_date:
            doc.add_paragraph(f'Дата рождения: {graduate.birth_date}')
        if graduate.city:
            doc.add_paragraph(f'Город: {graduate.city}')
        if graduate.phone:
            doc.add_paragraph(f'Телефон: {graduate.phone}')
        if hasattr(graduate, 'specialty') and graduate.specialty:
            doc.add_paragraph(f'Специальность: {graduate.specialty}')

        doc.add_paragraph()

        # Образование
        doc.add_heading('Образование', level=1)
        doc.add_paragraph(f'Университет: {graduate.university}')
        if graduate.program and graduate.program.name:
            doc.add_paragraph(f'Образовательная программа: {graduate.program.name}')
        elif graduate.program_master and graduate.program_master.name:
            doc.add_paragraph(f'Образовательная программа (магистратура): {graduate.program_master.name}')
        if graduate.grad_year:
            doc.add_paragraph(f'Год окончания: {graduate.grad_year}')

        doc.add_paragraph()

        # Опыт работы
        if graduate.experiences and len(graduate.experiences) > 0:
            doc.add_heading('Опыт работы', level=1)
            for exp in graduate.experiences:
                title_para = doc.add_paragraph()
                title_para.add_run(f'{exp.position}').bold = True
                doc.add_paragraph(f'Компания: {exp.company}', style='List Bullet')
                if exp.start_date:
                    end_date = exp.end_date if exp.end_date else 'настоящее время'
                    doc.add_paragraph(f'Период: {exp.start_date} — {end_date}', style='List Bullet')
                if exp.description:
                    doc.add_paragraph(f'Обязанности: {exp.description}', style='List Bullet')
                doc.add_paragraph()

        # Навыки
        if graduate.skills and len(graduate.skills) > 0:
            doc.add_heading('Профессиональные навыки', level=1)
            skills_text = ", ".join([s.name for s in graduate.skills])
            doc.add_paragraph(skills_text)
            doc.add_paragraph()

        # Сертификаты
        if graduate.certificates and len(graduate.certificates) > 0:
            doc.add_heading('Сертификаты и достижения', level=1)
            for cert in graduate.certificates:
                cert_text = cert.title
                if cert.issuer:
                    cert_text += f' ({cert.issuer})'
                if cert.issued_date:
                    cert_text += f' — {cert.issued_date}'
                doc.add_paragraph(cert_text, style='List Bullet')
            doc.add_paragraph()

        # Личные качества
        if graduate.personal_qualities:
            doc.add_heading('Личные качества', level=1)
            doc.add_paragraph(graduate.personal_qualities)

        # Дата составления
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(f'Дата составления: {datetime.now().strftime("%d.%m.%Y")}')

        # Создаём директорию если её нет
        os.makedirs("storage/cvs", exist_ok=True)

        # Формируем имя файла: Резюме_Фамилия.docx
        filename = f"Резюме_{clean_last_name}.docx"

        # Если файл с таким именем уже существует, добавляем дату
        filepath = os.path.join("storage", "cvs", filename)
        if os.path.exists(filepath):
            name_part = f"Резюме_{clean_last_name}"
            filename = f"{name_part}_{datetime.now().strftime('%Y%m%d')}.docx"
            filepath = os.path.join("storage", "cvs", filename)

        doc.save(filepath)

        return filename

    def _setup_styles(self, doc):
        """Настройка стилей документа"""
        try:
            # Настройка стиля для заголовков
            heading_style = doc.styles['Heading 1']
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            heading_style.font.name = 'Arial'

            # Настройка стиля для обычного текста
            normal_style = doc.styles['Normal']
            normal_style.font.size = Pt(12)
            normal_style.font.name = 'Times New Roman'

            # Настройка стиля для списков
            try:
                list_style = doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
                list_style.base_style = doc.styles['Normal']
                list_style.paragraph_format.left_indent = Inches(0.5)
                list_style.font.size = Pt(12)
            except:
                pass

        except Exception as e:
            print(f"Ошибка настройки стилей: {e}")